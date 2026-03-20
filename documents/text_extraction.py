"""
Extracción de texto de archivos (PDF, DOCX, XLS, XLSX) para indexación y búsqueda.
"""
import os
import logging

logger = logging.getLogger(__name__)


def extract_text_from_file(file_field):
    """
    Extrae texto de un archivo (FileField o path).
    Soporta: PDF, DOCX, XLS, XLSX.
    Retorna str o '' si no se puede extraer.
    """
    if not file_field:
        return ""
    path = getattr(file_field, "path", None) or (file_field if isinstance(file_field, (str, os.PathLike)) else None)
    if not path or not os.path.isfile(path):
        return ""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            return _extract_pdf(path)
        if ext in (".docx", ".doc"):
            return _extract_docx(path)
        if ext == ".xls":
            return _extract_xls(path)
        if ext == ".xlsx":
            return _extract_xlsx(path)
    except Exception as e:
        logger.warning("No se pudo extraer texto de %s: %s", path, e)
    return ""


def _extract_pdf(path):
    from pypdf import PdfReader
    reader = PdfReader(path)
    parts = []
    for page in reader.pages:
        try:
            t = page.extract_text()
            if t:
                parts.append(t)
        except Exception:
            pass
    return "\n".join(parts) if parts else ""


def _extract_docx(path):
    from docx import Document as DocxDocument
    doc = DocxDocument(path)
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts) if parts else ""


def _extract_xls(path):
    """Extrae texto de Excel .xls (Excel 97-2003) usando xlrd."""
    import xlrd

    def _norm(s: str) -> str:
        # Normaliza acentos comunes para detectar encabezados.
        s = (s or "").strip().lower()
        repl = {
            "á": "a", "é": "e", "í": "i", "ó": "o", "ú": "u",
            "Á": "a", "É": "e", "Í": "i", "Ó": "o", "Ú": "u",
            "ñ": "n", "Ñ": "n",
        }
        for k, v in repl.items():
            s = s.replace(k, v)
        return s

    book = xlrd.open_workbook(path)
    fallback_parts = []
    structured_parts = []
    qa_records = []
    max_rows_structured = 2000
    max_header_scan = 80

    target_cols = {
        "GC": None,
        "Tipo de Consulta": None,
        "Especialidad": None,
        "Pregunta": None,
        "Documento relacionado": None,
        "Respuesta": None,
        "Observación": None,
        "Respondido": None,
        "Columna1": None,
    }

    for sheet in book.sheets():
        # 1) Buscar encabezado (fila) y columnas por nombre (limitado a primeras filas)
        header_row_idx = None
        col_map = dict(target_cols)
        for row_idx in range(min(max_header_scan, sheet.nrows)):
            row_found_preg = False
            row_found_resp = False
            for col_idx in range(sheet.ncols):
                try:
                    cell = sheet.cell(row_idx, col_idx)
                    cell_val = cell.value
                    if cell_val is None:
                        continue
                    sval = _norm(str(cell_val))

                    if sval == "gc":
                        col_map["GC"] = col_idx
                    if sval == "tipo de consulta":
                        col_map["Tipo de Consulta"] = col_idx
                    if sval == "especialidad":
                        col_map["Especialidad"] = col_idx
                    if sval == "pregunta":
                        col_map["Pregunta"] = col_idx
                        row_found_preg = True
                    if sval == "documento relacionado":
                        col_map["Documento relacionado"] = col_idx
                    if sval == "respuesta":
                        col_map["Respuesta"] = col_idx
                        row_found_resp = True
                    if sval == "observacion":
                        col_map["Observación"] = col_idx
                    if sval == "respondido":
                        col_map["Respondido"] = col_idx
                    if sval == "columna1":
                        col_map["Columna1"] = col_idx
                except Exception:
                    continue

            # Solo tomamos esta fila como encabezado si están Pregunta y Respuesta
            if row_found_preg and row_found_resp:
                header_row_idx = row_idx
                break

        # 2) Si encontramos tabla con Pregunta/Respuesta, construir bloque estructurado
        def _cell_text(ridx: int, cidx: int):
            if cidx is None:
                return ""
            try:
                cell = sheet.cell(ridx, cidx)
                if cell.ctype == xlrd.XL_CELL_TEXT:
                    return str(cell.value or "").strip()
                if cell.ctype == xlrd.XL_CELL_NUMBER:
                    return str(cell.value).strip()
                if cell.ctype == xlrd.XL_CELL_DATE:
                    return str(cell.value).strip()
                # fallback
                return str(cell.value or "").strip()
            except Exception:
                return ""

        if header_row_idx is not None and col_map["Pregunta"] is not None and col_map["Respuesta"] is not None:
            count = 0
            for ridx in range(header_row_idx + 1, sheet.nrows):
                if count >= max_rows_structured:
                    break

                gc = _cell_text(ridx, col_map["GC"])
                tipo = _cell_text(ridx, col_map["Tipo de Consulta"])
                esp = _cell_text(ridx, col_map["Especialidad"])
                preg = _cell_text(ridx, col_map["Pregunta"])
                doc_rel = _cell_text(ridx, col_map["Documento relacionado"])
                resp = _cell_text(ridx, col_map["Respuesta"])
                obs = _cell_text(ridx, col_map["Observación"])
                respd = _cell_text(ridx, col_map["Respondido"])
                col1 = _cell_text(ridx, col_map["Columna1"])

                if not (preg or resp or gc or tipo or esp or doc_rel or obs or respd or col1):
                    continue

                preg_inner = (preg or "").strip().strip(" '\"")
                resp_inner = (resp or "").strip().strip(" '\"")
                if not preg_inner or not resp_inner:
                    continue

                # Normalizar para que:
                # - pregunta SIEMPRE sea: ¿...?
                # - respuesta SIEMPRE sea: @...@
                if preg_inner.startswith("¿"):
                    preg_inner = preg_inner[1:]
                if preg_inner.endswith("?"):
                    preg_inner = preg_inner[:-1]
                preg_inner = preg_inner.strip()
                if not preg_inner:
                    continue
                preg_full = "¿" + preg_inner + "?"

                if resp_inner.startswith("@"):
                    resp_inner = resp_inner[1:]
                if resp_inner.endswith("@"):
                    resp_inner = resp_inner[:-1]
                resp_inner = resp_inner.strip()
                if not resp_inner:
                    continue
                # Formato indexado simple pero con metadata para el modal:
                # - pregunta siempre entre ¿...?
                # - respuesta siempre entre @...@
                # - además guardamos GC y Especialidad
                qa_records.append(
                    "GC: %s\nEspecialidad: %s\nN: %s\n%s@%s@"
                    % (gc, esp, col1, preg_full, resp_inner)
                )
                count += 1
        # Fallback estructurado: si no se pudo formar qa_records (header raro / headers separados),
        # recorrer filas buscando directamente delimitadores ¿... ? y @...@ en las columnas detectadas.
        if not qa_records and col_map["Pregunta"] is not None and col_map["Respuesta"] is not None:
            count = 0
            for ridx in range(sheet.nrows):
                if count >= max_rows_structured:
                    break
                gc = _cell_text(ridx, col_map["GC"])
                esp = _cell_text(ridx, col_map["Especialidad"])
                preg = _cell_text(ridx, col_map["Pregunta"])
                resp = _cell_text(ridx, col_map["Respuesta"])
                if not preg or not resp:
                    continue
                if "¿" not in preg or "@" not in resp:
                    continue

                preg_inner = (preg or "").strip().strip(" '\"")
                resp_inner = (resp or "").strip().strip(" '\"")
                if not preg_inner or not resp_inner:
                    continue

                if preg_inner.startswith("¿"):
                    preg_inner = preg_inner[1:]
                if preg_inner.endswith("?"):
                    preg_inner = preg_inner[:-1]
                preg_inner = preg_inner.strip()
                if not preg_inner:
                    continue
                preg_full = "¿" + preg_inner + "?"

                if resp_inner.startswith("@"):
                    resp_inner = resp_inner[1:]
                if resp_inner.endswith("@"):
                    resp_inner = resp_inner[:-1]
                resp_inner = resp_inner.strip()
                if not resp_inner:
                    continue

                qa_records.append(
                    "GC: %s\nEspecialidad: %s\nN: %s\n%s@%s@"
                    % (gc, esp, col1, preg_full, resp_inner)
                )
                count += 1

        # 3) Fallback general: extraer todo lo demás (para no perder contenido)
        for row_idx in range(sheet.nrows):
            for col_idx in range(sheet.ncols):
                try:
                    cell = sheet.cell(row_idx, col_idx)
                    if cell.ctype == xlrd.XL_CELL_TEXT and str(cell.value).strip():
                        fallback_parts.append(str(cell.value))
                    elif cell.ctype == xlrd.XL_CELL_NUMBER:
                        fallback_parts.append(str(cell.value))
                    elif cell.ctype == xlrd.XL_CELL_DATE:
                        fallback_parts.append(str(cell.value))
                except Exception:
                    continue

    if qa_records:
        return "\n".join(qa_records)
    if qa_records:
        return "\n".join(qa_records)
    if qa_records:
        return "\n".join(qa_records)
    return "\n".join(fallback_parts) if fallback_parts else ""


def _extract_xlsx(path):
    """Extrae texto de Excel .xlsx usando openpyxl."""
    from openpyxl import load_workbook

    def _norm(s: str) -> str:
        s = (s or "").strip().lower()
        repl = {
            "á": "a", "é": "e", "í": "i", "ó": "o", "ú": "u",
            "Á": "a", "É": "e", "Í": "i", "Ó": "o", "Ú": "u",
            "ñ": "n", "Ñ": "n",
        }
        for k, v in repl.items():
            s = s.replace(k, v)
        return s

    wb = load_workbook(path, read_only=True, data_only=True)
    fallback_parts = []
    structured_parts = []
    qa_records = []
    max_rows_structured = 2000
    max_header_scan = 80

    target_cols = {
        "GC": None,
        "Tipo de Consulta": None,
        "Especialidad": None,
        "Pregunta": None,
        "Documento relacionado": None,
        "Respuesta": None,
        "Observación": None,
        "Respondido": None,
        "Columna1": None,
    }

    for sheet in wb.worksheets:
        header_row_num = None  # 1-based en openpyxl
        col_map = dict(target_cols)

        # 1) Detectar encabezados en las primeras filas
        row_found_preg = False
        row_found_resp = False
        for i, row in enumerate(
            sheet.iter_rows(min_row=1, max_row=max_header_scan, values_only=True),
            start=1
        ):
            row_found_preg = False
            row_found_resp = False
            for j, cell_val in enumerate(row):
                if cell_val is None:
                    continue
                sval = _norm(str(cell_val))

                # Headers: usar igualdad exacta para no capturar texto de datos (ej. "preguntar", "ver respuesta").
                if sval == "gc":
                    col_map["GC"] = j
                if sval == "tipo de consulta":
                    col_map["Tipo de Consulta"] = j
                if sval == "especialidad":
                    col_map["Especialidad"] = j
                if sval == "pregunta":
                    col_map["Pregunta"] = j
                    row_found_preg = True
                if sval == "documento relacionado":
                    col_map["Documento relacionado"] = j
                if sval == "respuesta":
                    col_map["Respuesta"] = j
                    row_found_resp = True
                if sval == "observacion":
                    col_map["Observación"] = j
                if sval == "respondido":
                    col_map["Respondido"] = j
                if sval == "columna1":
                    col_map["Columna1"] = j

            if row_found_preg and row_found_resp:
                header_row_num = i
                break

        # 2) Construir bloque estructurado por fila
        if header_row_num is not None and col_map["Pregunta"] is not None and col_map["Respuesta"] is not None:
            count = 0
            for ridx, row in enumerate(
                sheet.iter_rows(min_row=header_row_num + 1, max_row=sheet.max_row, values_only=True),
                start=header_row_num + 1
            ):
                if count >= max_rows_structured:
                    break

                values = list(row)

                def _v(cidx):
                    if cidx is None or cidx < 0 or cidx >= len(values):
                        return ""
                    v = values[cidx]
                    if v is None:
                        return ""
                    return str(v).strip()

                gc = _v(col_map["GC"])
                tipo = _v(col_map["Tipo de Consulta"])
                esp = _v(col_map["Especialidad"])
                preg = _v(col_map["Pregunta"])
                doc_rel = _v(col_map["Documento relacionado"])
                resp = _v(col_map["Respuesta"])
                obs = _v(col_map["Observación"])
                respd = _v(col_map["Respondido"])
                col1 = _v(col_map["Columna1"])

                if not (preg or resp or gc or tipo or esp or doc_rel or obs or respd or col1):
                    continue

                preg_inner = (preg or "").strip().strip(" '\"")
                resp_inner = (resp or "").strip().strip(" '\"")
                if not preg_inner or not resp_inner:
                    continue

                if preg_inner.startswith("¿"):
                    preg_inner = preg_inner[1:]
                if preg_inner.endswith("?"):
                    preg_inner = preg_inner[:-1]
                preg_inner = preg_inner.strip()
                if not preg_inner:
                    continue
                preg_full = "¿" + preg_inner + "?"

                if resp_inner.startswith("@"):
                    resp_inner = resp_inner[1:]
                if resp_inner.endswith("@"):
                    resp_inner = resp_inner[:-1]
                resp_inner = resp_inner.strip()
                if not resp_inner:
                    continue
                # Formato indexado simple pero con metadata para el modal:
                # - pregunta siempre entre ¿...?
                # - respuesta siempre entre @...@
                # - además guardamos GC, Especialidad y N (Columna1)
                qa_records.append(
                    "GC: %s\nEspecialidad: %s\nN: %s\n%s@%s@"
                    % (gc, esp, col1, preg_full, resp_inner)
                )
                count += 1

        # Fallback estructurado: si no se pudo formar qa_records (header raro / encabezados separados),
        # recorre todas las filas buscando directamente delimitadores ¿... ? y @...@ en las columnas detectadas.
        if not qa_records and col_map["Pregunta"] is not None and col_map["Respuesta"] is not None:
            count = 0
            for row in sheet.iter_rows(values_only=True):
                if count >= max_rows_structured:
                    break
                values = list(row)

                def _v(cidx):
                    if cidx is None or cidx < 0 or cidx >= len(values):
                        return ""
                    v = values[cidx]
                    if v is None:
                        return ""
                    return str(v).strip()

                gc = _v(col_map["GC"])
                esp = _v(col_map["Especialidad"])
                preg = _v(col_map["Pregunta"])
                resp = _v(col_map["Respuesta"])
                col1 = _v(col_map["Columna1"])

                if not preg or not resp:
                    continue
                if "¿" not in preg or "@" not in resp:
                    continue

                preg_inner = (preg or "").strip().strip(" '\"")
                resp_inner = (resp or "").strip().strip(" '\"")
                if not preg_inner or not resp_inner:
                    continue

                if preg_inner.startswith("¿"):
                    preg_inner = preg_inner[1:]
                if preg_inner.endswith("?"):
                    preg_inner = preg_inner[:-1]
                preg_inner = preg_inner.strip()
                if not preg_inner:
                    continue
                preg_full = "¿" + preg_inner + "?"

                if resp_inner.startswith("@"):
                    resp_inner = resp_inner[1:]
                if resp_inner.endswith("@"):
                    resp_inner = resp_inner[:-1]
                resp_inner = resp_inner.strip()
                if not resp_inner:
                    continue

                qa_records.append(
                    "GC: %s\nEspecialidad: %s\nN: %s\n%s@%s@"
                    % (gc, esp, col1, preg_full, resp_inner)
                )
                count += 1

        # 3) Fallback: extraer todo lo demás
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value is not None and str(cell.value).strip():
                    fallback_parts.append(str(cell.value))

    wb.close()

    if qa_records:
        return "\n".join(qa_records)
    return "\n".join(fallback_parts) if fallback_parts else ""
